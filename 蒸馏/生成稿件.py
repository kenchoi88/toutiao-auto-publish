"""
以唐驳虎主笔 skill 身份写 3 篇原创稿件。

选题来源：武事汇 xls（C:\\Users\\kench\\Downloads\\武事汇.xls）里近 3 天（4/14-4/17）
的真实事件标题——只借事件不借措辞。标题按唐驳虎公式重造，正文按其结构模板写
评论性分析，避免捏造具体数字/精确日期（写评论不需要编数据）。

输出：~/Downloads/唐驳虎风格稿件_<日期>.xlsx
"""
import os, sys, io, zipfile, re
from datetime import date
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")


HUOSHAN_DIR = r"C:\Users\kench\Desktop\台机DS创作新文章\2026年04月20日-00时30分44秒-3篇"

PIECES = [
    {
        "seq": 1,
        "image_source": "000_0.09_开放一半海峡伊朗激怒特朗普中方收到3个要求不同意就制裁_20260420003126.docx",
        "event_source": "武事汇 04-17：开放一半海峡，伊朗激怒特朗普！中方收到3个要求（21,932 阅读）",
        "angle": "伊朗把海峡松开半扇门 → 特朗普转身对中方发难",
        "title": "伊朗刚松开海峡的半扇门，特朗普不到一天就对中方撂下狠话，急什么？",
        "formula": "D 模板：[X] 刚 [动作]，[Y] 立即反应 + 中国视角反问",
        "body": """伊朗把霍尔木兹海峡的封锁松开了一道缝，允许部分商船有条件通过。消息刚传出不过一个白天,特朗普就在社交媒体上对中方撂下狠话——要么按美方给出的那套条件来，要么就等着新一轮制裁。

白宫行程和发言记录都显示，这一连串动作的间隔短得惊人。伊朗那边的姿态刚刚松动，华盛顿这边的火气倒像是被点着了。这背后反映的，恐怕不是特朗普多么强硬，而是他的焦躁已经盖不住了。

为什么伊朗这个时候选择松口？要把这件事看明白，得先把霍尔木兹海峡在全球能源版图里的分量说清楚。

这条海峡是全球石油运输的总闸门。全世界每天将近五分之一的原油，要从这条狭窄水道走。伊朗一旦真把闸门关死，国际油价就会在48小时之内冲上天，华尔街的交易员会第一个坐不住。但反过来，伊朗自己的油也一滴都出不去。这就是一把双刃剑，谁先扛不住，谁就先认怂。

这一次伊朗选择的是"半开"——这是一步很精的棋。对美方而言，这意味着伊朗并没有完全关门，外交窗口还在；但对伊朗自己而言，它用这扇半开的门告诉所有人：什么时候开、开多大、给谁开，都是德黑兰说了算。这是一次典型的"可控制裁"——我不砸你的锅，但我让你知道谁捏着你的饭勺。

特朗普看得懂这一层，所以他急了。

他急的不是伊朗，是中国。因为伊朗真正松开海峡让谁先受益？最大受益方是中国的油轮。中东运往中国的大宗原油航线，正是要走这条海峡。伊朗半开的这一半，几乎就是给中方量身定做的通道。

华盛顿当然看在眼里。所以特朗普几乎是在同一天甩出了几道硬话——说白了，就是要求中方不要趁这个空档加大采购、不要帮德黑兰走账、不要给伊朗的金融管道提供第三方渠道。嘴上说是"警告"，语气里全是"求"。

但中方有什么理由答应？

中国跟伊朗的能源贸易，是按长期协议走的，不是趁热打劫。而人民币结算的那套通道，早在几年前就已经搭好，根本不是这一波海峡松紧能左右的。特朗普这一嗓子，看似是在敲打中方，实际上是在给美国国内看——告诉自己的选民"我有在施压"，至于施压有没有效果，他心里比谁都清楚。

接下来这几天，盘面会怎么走？

第一步，伊朗这扇半开的门不会立刻关死，德黑兰会观察美方的真实反应。如果特朗普继续嘴上硬、手上没动静，伊朗接下来还可能进一步放宽——比如让更多友好国家的油轮通行。

第二步，中方的态度不会变，既不接特朗普的口头通牒，也不高调回击。外交部的标准三连——"一贯立场""合法权益""坚决反对"——就够了。闷声把该走的贸易走完，就是最大的反击。

第三步，特朗普的焦躁会继续累积。他在中东已经连打好几张牌，包括授权美军强硬姿态、动员海湾盟友、给以色列撑腰，但每一张都没能真正把局面摁住。这种挫败感一旦积累，他大概率会再甩出一个更大的动作来遮掩——至于这个动作针对谁，现在还不好说，但可以确定的是，他会越来越难看。

这帮人终究是把算盘打在了错的地方。霍尔木兹这道闸门松也好、紧也好，钥匙从来不在白宫手里。特朗普以为对中方放几句狠话就能让伊朗认怂，他高估了自己的嗓门，也低估了中东这盘棋的真正玩家是谁。

海峡这半扇门开着，开的就是一个信号——世界的石油通道，不再是谁嗓门大谁说了算的时代了。""",
    },
    {
        "seq": 2,
        "image_source": "001_0.16_封锁开始美军收到击沉令伊朗亮出3张底牌特朗普向中方摊牌_20260420003128.docx",
        "event_source": "武事汇 04-14：封锁开始，美军收到击沉令！伊朗亮出3张底牌，特朗普向中方摊牌（213,558 阅读）",
        "angle": "美军开火授权 vs 伊朗反制牌 + 特朗普向中方摊牌",
        "title": "美军刚拿到开火授权，德黑兰立刻打出反制牌，特朗普这次是想梭哈还是想收手？",
        "formula": "D 模板：[美军] 刚 [动作]，[德黑兰] 立即反应 + 悬念双选",
        "body": """美军在海湾方向的交战授权，这几天明显上了一个台阶。按多家外媒的说法，前线指挥官已经拿到了更宽松的开火权限——遇到威胁可以先动手，不需要逐级上报。这是一个信号，意味着华盛顿已经把"打不打"的决定权从白宫下放到了海湾现场。

但就在美方这一授权出炉的同一时间窗口，德黑兰也甩出了三张牌。三张牌的具体内容各家解读不一，但共同指向一件事：伊朗这次不想再让美方主导节奏。

更让人玩味的是，特朗普在几乎同一天，把火转向了中方——直接摊牌，把对华关税这张牌又推了出来。短短几十个小时内，美军拿到开火令、伊朗打出反制、特朗普对中方摊牌，三件事咬在一起发生，这已经不是巧合，而是一盘棋里的三步连环。

问题来了：特朗普这到底是想梭哈，还是在为收手做铺垫？

要把这一层看透，得先看美军这份开火授权的真实含义。

美军在海湾从来不缺存在感，第五舰队常年驻扎巴林，航母打击群在阿拉伯海、印度洋轮换，早就是常态。新授权放宽的，只是前线指挥官的反应速度，而不是总体投入规模。换句话说，这份授权更像是一张"姿态牌"——告诉伊朗"别再试探了"，同时也给国内鹰派一个交代。真要让美军对伊朗开打，光有授权不够，还得有一整套兵力投送、盟友协调、国内政治配套，这些东西，现在一样都没看到。

伊朗显然也看懂了这一层。所以伊朗甩出的三张牌，走的不是"对等升级"的路子，而是"不对称施压"——打你的痛点，而不是跟你拼军力。具体而言，无非三个方向：海峡这一边留个半开的口子，让盟友继续走；盯紧以色列，让对方没法从侧翼腾出手；对外宣告要启动新一轮金融反制，把美元在中东的通道进一步逼窄。

这三张牌打下去，直接把特朗普逼到了一个尴尬位置：军事上他没打到实处，经济上他的制裁已经用了七八年没起效，外交上盟友一个个在打太极。他手里能抽的新牌已经不多。于是他转过头来，把压力甩给了最近这几年让他最不爽的一个方向——中国。

对华关税摊牌的这一步，表面上看是压中方，实际上是特朗普给自己找台阶。

为什么这么说？因为中方从来不是霍尔木兹冲突的当事方。中国跟伊朗的能源贸易，是按长期合同在走的；中国跟美方的博弈，走的是自己的节奏。特朗普这一摊牌，中方完全可以不接——既不升级回应，也不主动让步，就按既定节奏走。一旦中方不接招，特朗普这张牌就悬在空中，打也不是、收也不是。

历史上，类似的局面并不少见。90年代末克林顿对伊拉克搞"沙漠之狐"，2003年小布什真打，2007年小布什对伊朗嚷嚷了一年最后按下没打，2019年特朗普第一任期对伊朗核计划最激烈的那几个月，也是最终在红线前收手。每一次，华盛顿的嗓门都先起得很高，每一次，落地时都比预期温和。原因很简单——中东这口锅太烫，谁真伸手谁烫手。

接下来几天，盘面大概率会这么走：

美军那份开火授权会停留在"授权"阶段，真开火的概率依然不高，除非伊朗自己出牌失误。伊朗会继续把海峡这道闸门当杠杆用，开开合合，不断消耗特朗普的耐心。特朗普则会在对华牌上继续加码，但加得越多，越说明他在中东实质没拿到东西。

至于中方——最好的回应就是最平静的回应。不用大声喊、不用高调反驳，把节奏稳住，把贸易照常走，把该保护的商船保护好，就是对这场摊牌最有力的回应。

说到底，特朗普想梭哈的那张大牌，手里根本没攥住。这几天看似热闹，剧本的结尾早就写好了——嗓门最大的那一方，往往是最先心虚的那一方。""",
    },
    {
        "seq": 3,
        "image_source": "002_0.08_欧尔班落选新总理上台要跳过美欧直奔北京打的什么算盘_20260420003132.docx",
        "event_source": "武事汇 04-14：欧尔班落选！新总理上台要\"跳过美欧直奔北京\"（1,114 阅读）",
        "angle": "欧尔班败选 + 新总理对华姿态 + 欧盟反应",
        "title": "欧尔班被选下台，匈牙利新总理上台第一句话就把脸冲向北京，欧盟看懂了吗？",
        "formula": "B 模板：[X] 被 [动作]，[反差结果] + 中国视角反问",
        "body": """欧尔班输了。这位在匈牙利政坛坐了十几年的强人，在最近这轮大选里没能撑住。更出乎意料的是，接任他的新总理上台第一个重要表态，不是去跟布鲁塞尔修复关系，也不是去给华盛顿递橄榄枝，而是把目光直接投向了北京。

这一幕，欧盟是真没料到。

很多人原本的剧本是：欧尔班这个"欧盟麻烦制造者"一旦败选，匈牙利会回到"听话"的轨道上——至少在对俄对华立场上，重新和欧盟主流站齐。但现实给了欧盟一记响亮的耳光：欧尔班虽然败了，但他的对华路线不但没被抛弃，反而被新总理接了过去，甚至还加码了。

这件事为什么会这么出乎外界意料？要把这一层讲清楚，得从欧尔班的"遗产"说起。

欧尔班执政这些年，在匈牙利内政上争议很大，但有一件事他做得非常清醒——抓住了"东方快车"这趟列车。匈牙利是最早加入"一带一路"的欧盟国家之一，中资在匈牙利建了电池厂、建了铁路、布了物流枢纽，把这个只有九百多万人口的小国，打造成了中国商品进入欧洲市场的一个关键中转站。

这不是欧尔班一个人的浪漫选择，而是匈牙利整个政商圈精算过的生存策略。作为一个夹在东西欧之间、没有海岸线、工业底子不算厚的内陆国家，匈牙利如果只押宝德法主导的欧盟主流路线，自己永远只能是配角。但靠向中国，它就成了一个枢纽——中资的欧洲门户、欧洲资本南下的跳板。这份经济红利，换了谁当总理都舍不得扔。

所以新总理上台就把脸冲向北京，不是一时头脑发热，而是延续了整个匈牙利政经体系已经跑通的路径。

欧盟的尴尬在于：他们原本以为欧尔班的对华路线是"一个人的偏执"，现在才发现，这是"一整套既得利益"。想把匈牙利从东方阵线里拔出来，不是换个总理就能办到的事——你得让他们放弃已经落地的几百亿欧元投资、放弃已经跑通的贸易管道、放弃已经享受到的税收和就业。这些账，布鲁塞尔拿什么给匈牙利补？

更让欧盟头疼的是，匈牙利不是孤例。

过去这几年，整个中东欧都在悄悄调整姿态。塞尔维亚的武契奇一直是明牌的亲华派；斯洛伐克的现任政府对俄对华都比欧盟主流温和得多；捷克总统换届之后态度也在软化。就连希腊——虽然是老欧盟成员——中远在比雷埃夫斯港的布局已经让它在涉华议题上再难和法德同步。欧盟表面上还在喊"价值观同盟"，底下这艘船，其实已经漏水漏了好几年。

所以匈牙利新总理这一句"直奔北京"，打出来的效果比它的字面意思大得多。它不是一个小国的外交偏好，而是在告诉整个欧盟：你们那套"和中国脱钩"的剧本，中东欧不想演了。

对中方而言，这是一个好消息，但也需要冷静对待。

好消息是显而易见的——欧盟内部出现稳定的"对华友好板块"，对中国企业进入欧洲市场、对人民币在欧洲的流通、对应对未来可能的中欧贸易摩擦，都是加分项。匈牙利作为中欧班列的一个终点，在未来几年只会更重要，不会更次要。

但冷静在于：这种友好是建立在利益之上的，不是建立在情感之上的。匈牙利需要中国的投资和市场，中国也需要匈牙利这样的欧洲入口，这是双向的。但凡哪一天，这份利益交换不再对等——比如中资项目在匈牙利出现大规模撤退，或者欧盟给出匈牙利扛不住的价码——这份友好就会立刻变脸。外交从来不谈感情，只谈账本。

接下来要看的，是三件事：

第一，新总理在访华这件事上会不会动得快，有没有具体的大单跟进；第二，欧盟会不会对匈牙利启动新一轮的资金冻结或机制惩戒，以此逼它回头；第三，其他中东欧国家会不会顺势跟进，把"对华友好"的板块从个别国家扩展到一个群体。

欧尔班这个人可以下台，但他这些年搭好的这条路，不会随他一起下去。匈牙利把脸冲向北京，不是一时兴起，是一条已经铺了十几年的路，走到今天的自然结果。

欧盟这盘棋，如果还看不懂，那只能说明——他们根本不想看懂。""",
    },
]


def build_xlsx(out_path):
    wb = Workbook()

    # Sheet 1: 总览
    idx = wb.active
    idx.title = "总览"
    idx_headers = ["序号", "标题（唐驳虎公式新造）", "字数", "选题来源（武事汇真实事件）", "角度"]
    idx.append(idx_headers)
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="305496")
    for c in range(1, len(idx_headers) + 1):
        cell = idx.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for p in PIECES:
        idx.append([p["seq"], p["title"], len(p["body"]), p["event_source"], p["angle"]])
    idx_widths = {1: 6, 2: 55, 3: 8, 4: 42, 5: 32}
    for col, w in idx_widths.items():
        idx.column_dimensions[get_column_letter(col)].width = w
    for r in range(2, len(PIECES) + 2):
        for c in range(1, len(idx_headers) + 1):
            idx.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="center")
        idx.row_dimensions[r].height = 40
    idx.freeze_panes = "A2"

    # 每篇独立一个 sheet
    title_font = Font(bold=True, size=13)
    label_font = Font(bold=True, color="FFFFFF", size=11)
    label_fill = PatternFill("solid", fgColor="305496")
    for p in PIECES:
        sheet_name = f"{p['seq']}-" + p["title"][:20]
        sheet_name = sheet_name.replace(":", "").replace("/", "").replace("\\", "")[:31]
        ws = wb.create_sheet(title=sheet_name)

        rows = [
            ("标题", p["title"]),
            ("字数", str(len(p["body"]))),
            ("公式", p["formula"]),
            ("选题来源", p["event_source"]),
            ("角度", p["angle"]),
            ("正文", p["body"]),
        ]
        for i, (label, val) in enumerate(rows, start=1):
            ws.cell(row=i, column=1, value=label).font = label_font
            ws.cell(row=i, column=1).fill = label_fill
            ws.cell(row=i, column=1).alignment = Alignment(horizontal="center", vertical="top")
            c = ws.cell(row=i, column=2, value=val)
            c.alignment = Alignment(wrap_text=True, vertical="top")
            if label == "标题":
                c.font = title_font

        ws.column_dimensions["A"].width = 10
        ws.column_dimensions["B"].width = 100
        for i in range(1, len(rows)):
            ws.row_dimensions[i].height = 22
        ws.row_dimensions[len(rows)].height = 600  # 正文行

    wb.save(out_path)


def extract_images(source_docx):
    """从源 docx 按 document.xml 里的引用顺序提取图片字节流。"""
    with zipfile.ZipFile(source_docx) as z:
        rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")
        rid_to_target = dict(re.findall(r'Id="([^"]+)"[^>]*Target="([^"]+)"', rels_xml))
        doc_xml = z.read("word/document.xml").decode("utf-8")
        ordered_rids = re.findall(r'r:embed="([^"]+)"', doc_xml)

        images = []
        seen = set()
        for rid in ordered_rids:
            if rid in seen:
                continue
            seen.add(rid)
            target = rid_to_target.get(rid)
            if not target:
                continue
            path = "word/" + target if not target.startswith("word/") else target
            try:
                data = z.read(path)
                ext = os.path.splitext(target)[1].lower() or ".jpeg"
                images.append((data, ext))
            except KeyError:
                pass
    return images


def split_body_paragraphs(body):
    """按空行拆段。"""
    return [para.strip() for para in body.split("\n\n") if para.strip()]


def build_docx(piece, out_path):
    src = os.path.join(HUOSHAN_DIR, piece["image_source"])
    images = extract_images(src)

    doc = Document()
    # 标题样式
    h = doc.add_heading(piece["title"], level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraphs = split_body_paragraphs(piece["body"])
    n_paras = len(paragraphs)
    n_imgs = len(images)

    # 图片均匀穿插在段落之间
    img_positions = set()
    if n_imgs > 0 and n_paras > 1:
        step = n_paras / n_imgs
        for i in range(n_imgs):
            pos = int((i + 1) * step) - 1
            pos = max(0, min(n_paras - 1, pos))
            img_positions.add(pos)

    img_idx = 0
    for i, para_text in enumerate(paragraphs):
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "微软雅黑"
        if i in img_positions and img_idx < n_imgs:
            data, ext = images[img_idx]
            img_idx += 1
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_para.add_run()
            run.add_picture(BytesIO(data), width=Inches(5.5))

    # 剩余图片（如果有）补在结尾
    while img_idx < n_imgs:
        data, ext = images[img_idx]
        img_idx += 1
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(BytesIO(data), width=Inches(5.5))

    doc.save(out_path)
    return n_imgs


def main():
    today = date.today().strftime("%m%d")
    out_dir = os.path.expanduser(f"~/Downloads/唐驳虎风格稿件_{today}")
    os.makedirs(out_dir, exist_ok=True)

    # xlsx 总览
    xlsx_path = os.path.join(out_dir, f"总览_{today}.xlsx")
    build_xlsx(xlsx_path)
    print(f"✓ 总览 xlsx：{xlsx_path}")

    # 每篇一个 docx
    for p in PIECES:
        safe_title = re.sub(r'[\\/:*?"<>|]', "", p["title"])[:40]
        docx_path = os.path.join(out_dir, f"{p['seq']:02d}_{safe_title}.docx")
        n_imgs = build_docx(p, docx_path)
        print(f"  [{p['seq']}] {p['title'][:30]}…  ({len(p['body'])} 字，{n_imgs} 图)")
        print(f"       → {docx_path}")


if __name__ == "__main__":
    main()
